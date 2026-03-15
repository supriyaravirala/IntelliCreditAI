from docx import Document
from datetime import datetime


def generate_cam(score, level, borrower, financials, research):

    doc = Document()

    # TITLE
    doc.add_heading('CREDIT APPRAISAL MEMORANDUM (CAM)', 0)
    doc.add_paragraph("Confidential | " + datetime.now().strftime("%d %B %Y"))

    doc.add_paragraph("")

    # -------------------------------------------------
    # 1. LOAN APPLICATION SUMMARY
    # -------------------------------------------------

    doc.add_heading("1. LOAN APPLICATION SUMMARY", level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    headers = ["Field", "Details"]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    summary = {
        "Borrower Name": borrower,
        "Date of Appraisal": datetime.now().strftime("%d-%m-%Y"),
        "Credit Score": f"{score} / 100",
        "Grade": "A" if score > 80 else "B",
        "Decision": "APPROVE" if score > 70 else "REVIEW",
        "Suggested Loan Limit": "INR 0 Crore",
        "Indicative Rate": "Base + 0.75%",
        "Web Research Risk": level
    }

    for key, value in summary.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)

    # -------------------------------------------------
    # 2. KEY FINANCIALS
    # -------------------------------------------------

    doc.add_paragraph("")
    doc.add_heading("2. KEY FINANCIALS", level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"

    for key, value in financials.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)

    # -------------------------------------------------
    # 3. RED FLAGS
    # -------------------------------------------------

    doc.add_paragraph("")
    doc.add_heading("3. RED FLAGS & RISK SIGNALS", level=1)

    doc.add_paragraph("⚠ Could not parse financial figures clearly", style='List Bullet')

    # -------------------------------------------------
    # 4. SECONDARY RESEARCH
    # -------------------------------------------------

    doc.add_paragraph("")
    doc.add_heading("4. SECONDARY RESEARCH FINDINGS", level=1)

    doc.add_paragraph(research)

    # -------------------------------------------------
    # 5. CREDIT ASSESSMENT
    # -------------------------------------------------

    doc.add_paragraph("")
    doc.add_heading("5. DETAILED CREDIT ASSESSMENT", level=1)

    doc.add_paragraph(
        "The borrower demonstrates stable financial performance with moderate leverage. "
        "However, limited financial visibility and missing EBITDA data create uncertainty "
        "around cash flow stability."
    )

    doc.add_paragraph(
        "Based on the AI credit model, the borrower has a credit score of "
        + str(score) + "/100 which indicates a " + level + " risk level."
    )

    # -------------------------------------------------
    # 6. FINAL RECOMMENDATION
    # -------------------------------------------------

    doc.add_paragraph("")
    doc.add_heading("6. FINAL RECOMMENDATION", level=1)

    decision = "APPROVE" if score > 70 else "REVIEW"

    doc.add_paragraph(
        f"DECISION: {decision} | LIMIT: INR 0 Crore | RATE: Base + 0.75%"
    )

    doc.add_paragraph("")
    doc.add_paragraph("Prepared by: Intelli-Credit AI Engine")
    doc.add_paragraph("Generated on: " + datetime.now().strftime("%d-%m-%Y %H:%M"))

    file_name = "CAM_Report.docx"

    doc.save(file_name)

    return file_name